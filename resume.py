import streamlit as st

from docx import Document
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate

from config import DEFAULT_RESUME_MODEL
from ollama_client import OllamaError, generate_text, list_models

st.set_page_config(
    page_title="AI Resume Builder",
    layout="wide",
)

st.title("AI Resume Builder")

st.sidebar.title("Settings")

try:
    available_models = list_models()
except OllamaError as exc:
    st.error(str(exc))
    st.stop()

if not available_models:
    st.error("No Ollama models found. Install one with `ollama pull llama3.2` or `ollama pull gemma3`.")
    st.stop()

default_model_index = (
    available_models.index(DEFAULT_RESUME_MODEL)
    if DEFAULT_RESUME_MODEL in available_models
    else 0
)

model_name = st.sidebar.selectbox(
    "Model",
    available_models,
    index=default_model_index,
)

resume_type = st.sidebar.selectbox(
    "Resume Type",
    [
        "Fresher",
        "Experienced",
        "ATS Friendly",
    ],
)

photo = st.file_uploader(
    "Upload Profile Photo (Optional)",
    type=["jpg", "jpeg", "png"],
)

with st.form("resume"):
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone")
    address = st.text_area("Address")
    objective = st.text_area("Career Objective")
    education = st.text_area("Education")
    skills = st.text_area("Skills")
    projects = st.text_area("Projects")
    experience = st.text_area("Experience")
    certifications = st.text_area("Certifications")
    achievements = st.text_area("Achievements")
    languages = st.text_input("Languages")
    generate = st.form_submit_button("Generate Resume")

if generate:
    prompt = f"""
Create a professional {resume_type} resume.

Name:
{name}

Email:
{email}

Phone:
{phone}

Address:
{address}

Career Objective:
{objective}

Education:
{education}

Skills:
{skills}

Projects:
{projects}

Experience:
{experience}

Certifications:
{certifications}

Achievements:
{achievements}

Languages:
{languages}

Use proper headings.
Use bullet points.
Return only the resume.
"""

    with st.spinner("Generating Resume..."):
        try:
            resume = generate_text(
                prompt,
                model=model_name,
                temperature=0.4,
                max_tokens=1200,
            )
        except OllamaError as exc:
            st.error(str(exc))
            st.stop()

    st.success("Resume Generated Successfully")

    if photo:
        st.image(photo, width=150)

    st.markdown(resume)

    st.download_button(
        "Download TXT",
        resume,
        "resume.txt",
        "text/plain",
    )

    doc = Document()
    doc.add_heading("Resume", level=1)
    doc.add_paragraph(resume)
    doc.save("resume.docx")

    with open("resume.docx", "rb") as file:
        st.download_button(
            "Download Word",
            file,
            "resume.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate("resume.pdf")
    story = []

    for line in resume.split("\n"):
        story.append(Paragraph(line, styles["BodyText"]))

    pdf.build(story)

    with open("resume.pdf", "rb") as file:
        st.download_button(
            "Download PDF",
            file,
            "resume.pdf",
            "application/pdf",
        )
